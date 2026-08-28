"""Word (.docx) rendering of the two Markdown deliverables.

WHY THIS EXISTS
---------------
`gtm_plan.md` and `executive_summary.md` are the artifacts people actually
forward — to a CMO, a sales lead, an agency. Markdown renders as raw syntax
in Outlook, Word, and most document viewers, so forwarding one means the
recipient sees `## Top 3 GTM Plays` instead of a heading. The .docx is the
same content in a format anyone can open, comment on, and pass along.

The Markdown these files contain is a small, known subset: ATX headings,
bullet and numbered lists, blockquotes, paragraphs, and `**bold**`. This
renders that subset faithfully and treats anything else as body text rather
than guessing — an unrecognized construct should look plain, never wrong.

The executive summary is not Markdown at all; it uses bare uppercase section
labels (``WHAT CHANGED``). Those are promoted to headings too, so both
documents come out looking like documents.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

log = logging.getLogger("demand_radar.docx")

#: ``## `` -> Heading 2, and so on. ``# `` is the document title.
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^\s*[-*]\s+(.*)$")
_NUMBERED = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_QUOTE = re.compile(r"^>\s?(.*)$")
_BOLD = re.compile(r"\*\*(.+?)\*\*")

#: An all-caps line short enough to be a label, not a sentence. Used for the
#: executive summary, which has no Markdown markup to go on.
_MAX_BARE_HEADING_LEN = 60


class DocxUnavailable(RuntimeError):
    """python-docx is not importable in this environment."""


def _add_runs(paragraph, text: str) -> None:
    """Split on ``**bold**`` so emphasis survives into Word."""
    for i, part in enumerate(_BOLD.split(text)):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = i % 2 == 1  # odd segments are the captured bold groups


def _is_bare_heading(line: str) -> bool:
    stripped = line.strip()
    return (
        bool(stripped)
        and len(stripped) <= _MAX_BARE_HEADING_LEN
        and any(ch.isalpha() for ch in stripped)
        and stripped == stripped.upper()
        and not stripped.endswith((".", ":", ","))
    )


def markdown_to_docx(text: str, path: Path, fallback_title: str) -> Path:
    """Render `text` to a .docx at `path`; returns the path written.

    `fallback_title` titles the document when the source opens with no `# `
    heading, which is the case for the executive summary.
    """
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - depends on install state
        raise DocxUnavailable(
            "python-docx is not installed, so no Word version was written. "
            "Install it with `pip install python-docx` (or reinstall the "
            "package with `pip install -e .`). The Markdown file is complete "
            "and unaffected."
        ) from exc

    lines = text.splitlines()
    title, body = fallback_title, lines

    # A title is always written, either from a leading `# ` or from
    # `fallback_title`, so `##` is always the document's top section and is
    # demoted to Heading 1 either way. Keying this off the `# ` alone left
    # a document that relies on the fallback — the executive summary —
    # starting at Heading 2 with Heading 1 unused.
    demoted = True

    # Only a leading `# ` counts as the title. Consuming it here — rather
    # than mid-loop — means the title is always the first thing in the
    # document, which is what Word's navigation pane and file previews read.
    for i, line in enumerate(lines):
        if not line.strip():
            continue
        if (match := _HEADING.match(line.strip())) and len(match.group(1)) == 1:
            title = match.group(2).strip()
            body = lines[i + 1 :]
        break

    doc = Document()
    _add_runs(doc.add_heading(level=0), title)

    # Markdown treats a single newline inside a block as a soft break, so
    # consecutive lines are buffered and joined rather than emitted one Word
    # paragraph each. Without this, prose wrapped at ~72 columns — which is
    # how the LLM prompts ask for it — arrives as separately-spaced sentence
    # fragments instead of paragraphs.
    pending: list[str] = []
    pending_style: str | None = None

    def flush() -> None:
        nonlocal pending, pending_style
        if pending:
            _add_runs(doc.add_paragraph(style=pending_style), " ".join(pending))
        pending = []
        pending_style = None

    for line in body:
        stripped = line.strip()
        if not stripped:
            flush()
            continue

        if match := _HEADING.match(stripped):
            flush()
            level = len(match.group(1)) - (1 if demoted else 0)
            _add_runs(
                doc.add_heading(level=min(max(level, 1), 4)), match.group(2).strip()
            )
            continue

        if match := _QUOTE.match(stripped):
            if pending_style != "Intense Quote":
                flush()
                pending_style = "Intense Quote"
            pending.append(match.group(1).strip())
            continue

        if match := _BULLET.match(line):
            flush()
            pending_style = "List Bullet"
            pending.append(match.group(1).strip())
            continue

        if match := _NUMBERED.match(line):
            flush()
            pending_style = "List Number"
            pending.append(match.group(1).strip())
            continue

        if _is_bare_heading(stripped):
            flush()
            _add_runs(doc.add_heading(level=1), stripped)
            continue

        # Plain text. A quote block ends here, since an unprefixed line is no
        # longer quoted. After a list item it is that item's continuation,
        # which is what Markdown's lazy continuation does.
        if pending_style == "Intense Quote":
            flush()
        pending.append(stripped)

    flush()

    doc.save(str(path))
    return path
