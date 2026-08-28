"""End-to-end pipeline behavior with mock providers, plus failure handling."""

import json
import zipfile

import pytest
from docx import Document

from demand_radar.demo import demo_config
from demand_radar.pipeline import Pipeline
from demand_radar.reporting import ConsoleReporter, RecordingReporter
from demand_radar.providers.llm.base import LLMProvider
from demand_radar.providers.llm.mock import MockLLMProvider
from demand_radar.providers.llm.router import LLMRouter, build_router
from demand_radar.providers.search.base import SearchError, SearchProvider
from demand_radar.providers.search.mock import MockSearchProvider
from demand_radar.schemas.evidence import EvidenceRow
from demand_radar.schemas.queries import QuerySet


def test_mock_search_is_deterministic():
    p = MockSearchProvider()
    assert p.search("enterprise voice AI", limit=5) == p.search(
        "enterprise voice AI", limit=5
    )
    assert len(p.search("x", limit=7)) == 7


def test_mock_search_urls_are_stable_across_processes():
    """Repeated calls within one process reuse the same hash salt, so they
    can't catch a slug derived from Python's per-process-salted hash() —
    only a pinned literal value catches that regression, since the old
    implementation would print a *different* value on every fresh
    interpreter run (e.g. every CI job or every `demand-radar demo`)."""
    urls = [r["url"] for r in MockSearchProvider().search("enterprise voice AI", limit=3)]
    assert urls == [
        "https://cx-operations-review.example.com/articles/82462",
        "https://cx-operations-review.example.com/articles/19930",
        "https://market-signal-daily.example.com/articles/36832",
    ]


def test_mock_llm_query_expansion_schema():
    out = MockLLMProvider().complete("query_expansion", "prompt", schema=QuerySet)
    assert isinstance(out, QuerySet)
    assert out.total() > 0


def test_full_mock_pipeline_writes_all_artifacts(tmp_path):
    cfg = demo_config()
    router = build_router(cfg.llm)
    pipe = Pipeline(cfg, router, MockSearchProvider(), output_dir=tmp_path, reporter=ConsoleReporter(echo=False))
    summary = pipe.run()
    assert "SYNTHETIC" in summary
    for name in (
        "queries.json",
        "evidence.json",
        "evidence.csv",
        "signals.json",
        "analysis.json",
        "gtm_plan.md",
        "executive_summary.md",
        "run_metadata.json",
    ):
        assert (tmp_path / name).exists(), name

    meta = json.loads((tmp_path / "run_metadata.json").read_text())
    assert meta["brand"] == "ElevenLabs"
    assert meta["normalized_evidence_rows"] > 0
    assert meta["llm_provider"] == "mock"


def test_analysis_evidence_ids_reference_real_rows(tmp_path):
    cfg = demo_config()
    router = build_router(cfg.llm)
    pipe = Pipeline(cfg, router, MockSearchProvider(), output_dir=tmp_path, reporter=ConsoleReporter(echo=False))
    pipe.run()
    evidence = json.loads((tmp_path / "evidence.json").read_text())
    valid_ids = {row["evidence_id"] for row in evidence}
    analysis = json.loads((tmp_path / "analysis.json").read_text())
    referenced = [
        eid
        for trend in analysis["trends"]
        for eid in trend["supporting_evidence_ids"]
    ] + [eid for sig in analysis["buying_signals"] for eid in sig["evidence_ids"]]
    assert referenced, "mock analysis should reference some evidence"
    assert set(referenced) <= valid_ids


class _BadJSONProvider(LLMProvider):
    """Always returns non-JSON garbage, to exercise failure handling."""

    name = "bad"

    def complete(self, task, prompt, model, schema=None, reasoning_level="medium"):
        if schema is None:
            return "not json"
        raise ValueError("malformed JSON from model")


def test_invalid_llm_response_surfaces_cleanly(tmp_path):
    cfg = demo_config()
    router = LLMRouter(_BadJSONProvider(), cfg.llm)
    pipe = Pipeline(cfg, router, MockSearchProvider(), output_dir=tmp_path, reporter=ConsoleReporter(echo=False))
    with pytest.raises(ValueError, match="malformed JSON"):
        pipe.run()


class _EmptySearchProvider(SearchProvider):
    """Every query returns zero results, as if credentials had expired."""

    name = "empty"

    def search(self, query, limit=10):
        return []


def test_run_aborts_when_no_evidence_collected(tmp_path):
    """A silent zero-evidence run would print a confident-looking report
    with nothing behind it — the pipeline must fail loudly instead."""
    cfg = demo_config()
    router = build_router(cfg.llm)
    pipe = Pipeline(cfg, router, _EmptySearchProvider(), output_dir=tmp_path, reporter=ConsoleReporter(echo=False))
    with pytest.raises(SearchError, match="No evidence collected"):
        pipe.run()
    assert not (tmp_path / "gtm_plan.md").exists()


def test_run_clears_stale_downstream_artifacts_before_writing(tmp_path):
    """A prior run's analysis/plan must not survive next to a failed run's
    fresh (but incomplete) evidence — that would look like one coherent,
    valid snapshot when it's actually two runs mixed together."""
    (tmp_path / "gtm_plan.md").write_text("STALE PLAN FROM A PRIOR RUN")
    (tmp_path / "run_metadata.json").write_text("STALE METADATA")

    cfg = demo_config()
    router = build_router(cfg.llm)
    pipe = Pipeline(cfg, router, MockSearchProvider(), output_dir=tmp_path, reporter=ConsoleReporter(echo=False))
    pipe.run()

    assert "STALE" not in (tmp_path / "gtm_plan.md").read_text()
    assert "STALE" not in (tmp_path / "run_metadata.json").read_text()


def _run_with_themes(tmp_path, themes_yaml) -> RecordingReporter:
    """Run the mock pipeline with a given taxonomy; return what it reported."""
    themes = tmp_path / "themes.yaml"
    themes.write_text(themes_yaml, encoding="utf-8")
    cfg = demo_config().model_copy(update={"themes_file": str(themes)})
    reporter = RecordingReporter()
    Pipeline(
        cfg,
        build_router(cfg.llm),
        MockSearchProvider(),
        output_dir=tmp_path / "out",
        reporter=reporter,
    ).run()
    return reporter


def test_warns_when_taxonomy_does_not_fit_the_evidence(tmp_path):
    """A mismatched taxonomy still produces counts — the warning is the only
    thing standing between a user and meaningless numbers under a correct
    brand header.

    Asserted as an event, not as a sentence: the finding is that coverage
    was 0%, and rewording the message should not break this test.
    """
    reported = _run_with_themes(
        tmp_path, "themes:\n  agriculture:\n    - tractor\n    - soybean\n"
    )
    assert len(reported.coverage_warnings) == 1
    warning = reported.coverage_warnings[0]
    assert warning.coverage == 0.0
    assert warning.brand == "ElevenLabs"
    assert warning.source.endswith("themes.yaml")


def test_no_warning_when_taxonomy_fits(tmp_path):
    reported = _run_with_themes(
        tmp_path, "themes:\n  broad:\n    - a\n    - e\n    - i\n"
    )
    assert reported.coverage_warnings == []


def test_run_writes_word_versions_of_both_deliverables(tmp_path):
    """The .docx twins are what actually get forwarded to other people."""
    cfg = demo_config()
    Pipeline(cfg, build_router(cfg.llm), MockSearchProvider(), output_dir=tmp_path).run()

    for stem in ("gtm_plan", "executive_summary"):
        docx = tmp_path / f"{stem}.docx"
        assert docx.exists(), f"{stem}.docx missing"
        assert zipfile.is_zipfile(docx), f"{stem}.docx is not a valid package"
        paragraphs = Document(str(docx)).paragraphs
        assert paragraphs[0].style.name == "Title", f"{stem}.docx has no title"
        # The Word version must carry the same content as its Markdown twin,
        # not just be a well-formed empty document.
        body = "\n".join(p.text for p in paragraphs[1:])
        markdown = (tmp_path / f"{stem}.md").read_text(encoding="utf-8")
        for line in [ln.strip() for ln in markdown.splitlines() if ln.strip()][-3:]:
            assert line.lstrip("#>-0123456789. ") in body

        # Every `## ` section must arrive as a Word heading. The content
        # check above cannot see the difference: a document rendered as one
        # unbroken block of body text is still a valid .docx and still
        # contains every sentence. It is just unreadable as a document,
        # which is the whole reason the .docx exists.
        sections = [ln[3:].strip() for ln in markdown.splitlines() if ln.startswith("## ")]
        assert sections, f"{stem}.md has no `## ` sections"
        headings = {
            p.text for p in paragraphs if p.style.name.startswith("Heading")
        }
        assert not [s for s in sections if s not in headings], (
            f"{stem}.docx renders its sections as body text: "
            f"{[s for s in sections if s not in headings]}"
        )

    # The summary has no `# ` heading of its own, so its title is the one the
    # pipeline supplies — the brand's name, which is what a recipient sees.
    summary_title = Document(str(tmp_path / "executive_summary.docx")).paragraphs[0]
    assert cfg.brand_name in summary_title.text


def test_stale_word_files_are_cleared_like_their_markdown(tmp_path):
    """The .docx files are run artifacts too. If they survived a failed run
    they would be the ones forwarded — stale content under a fresh date."""
    tmp_path.mkdir(parents=True, exist_ok=True)
    (tmp_path / "gtm_plan.docx").write_text("STALE", encoding="utf-8")

    cfg = demo_config()
    Pipeline(cfg, build_router(cfg.llm), MockSearchProvider(), output_dir=tmp_path).run()

    assert zipfile.is_zipfile(tmp_path / "gtm_plan.docx")


def test_run_survives_a_word_rendering_failure(tmp_path, monkeypatch):
    """The .md is the source of truth and the LLM calls are already paid for,
    so a Word problem must degrade to a reported finding, not fail the run."""
    import demand_radar.run_ledger as ledger_mod
    from demand_radar.docx_export import DocxUnavailable

    def _boom(*args, **kwargs):
        raise DocxUnavailable("python-docx is not installed")

    # The renderer is invoked by the ledger, which owns the Word rendition.
    monkeypatch.setattr(ledger_mod, "markdown_to_docx", _boom)

    cfg = demo_config()
    reported = RecordingReporter()
    Pipeline(
        cfg,
        build_router(cfg.llm),
        MockSearchProvider(),
        output_dir=tmp_path,
        reporter=reported,
    ).run()

    assert (tmp_path / "gtm_plan.md").read_text(encoding="utf-8").strip()
    assert not (tmp_path / "gtm_plan.docx").exists()
    assert {d.path.name for d in reported.degradations} == {
        "gtm_plan.docx",
        "executive_summary.docx",
    }
    assert all(isinstance(d.error, DocxUnavailable) for d in reported.degradations)
    # Nothing may point at a Word file that was never written.
    assert not reported.completions[0].manifest.wrote("gtm_plan", ".docx")


def test_a_completed_run_reports_the_word_file_it_wrote(tmp_path):
    cfg = demo_config()
    reported = RecordingReporter()
    Pipeline(
        cfg,
        build_router(cfg.llm),
        MockSearchProvider(),
        output_dir=tmp_path,
        reporter=reported,
    ).run()

    manifest = reported.completions[0].manifest
    assert manifest.wrote("gtm_plan", ".docx")
    assert manifest.path("gtm_plan", ".docx") == tmp_path / "gtm_plan.docx"


def test_analyze_reports_no_artifacts_it_does_not_write(tmp_path):
    """`analyze` replays stages 5-8 over evidence passed in with --input, so
    it never writes evidence.csv. Reporting it would send someone looking
    for a file that was never there."""
    cfg = demo_config()
    source = tmp_path / "src"
    Pipeline(cfg, build_router(cfg.llm), MockSearchProvider(), output_dir=source).run()
    rows = [EvidenceRow(**r) for r in json.loads((source / "evidence.json").read_text())]

    out = tmp_path / "replay"
    reported = RecordingReporter()
    Pipeline(
        cfg,
        build_router(cfg.llm),
        MockSearchProvider(),
        output_dir=out,
        reporter=reported,
    ).analyze_only(rows)

    assert not (out / "evidence.csv").exists()
    assert not reported.completions[0].manifest.wrote("evidence", ".csv")
    assert (out / "gtm_plan.md").exists()
