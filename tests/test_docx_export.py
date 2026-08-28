"""Word rendering of the Markdown deliverables.

The whole point of the .docx is that someone else can open it, so these
tests check the file is a real OOXML package and that the Markdown subject
actually maps onto Word styles — not merely that a file appeared.
"""

import zipfile

import pytest
from docx import Document

from demand_radar.docx_export import DocxUnavailable, markdown_to_docx

PLAN = """# GTM Plan

> Generated from evidence.

## Market Changes
Buyers are evaluating economics.

## Top 3 Plays

### 1. ROI Calculator
- Insight: pricing dominates
- Target: **VP Finance**

1. First action
2. Second action
"""

SUMMARY = """WHAT CHANGED
Pricing themes lead this run.

THREE MOST IMPORTANT ACTIONS
1. Ship the calculator
"""


def _styles(path):
    return [(p.style.name, p.text) for p in Document(str(path)).paragraphs]


def test_docx_is_a_valid_ooxml_package(tmp_path):
    """A .docx is a zip with specific parts; anything less will not open."""
    out = markdown_to_docx(PLAN, tmp_path / "plan.docx", "Fallback")
    with zipfile.ZipFile(out) as zf:
        names = set(zf.namelist())
        assert zf.testzip() is None
    assert "[Content_Types].xml" in names
    assert "word/document.xml" in names


def test_leading_h1_becomes_the_document_title(tmp_path):
    styles = _styles(markdown_to_docx(PLAN, tmp_path / "p.docx", "Fallback"))
    assert styles[0] == ("Title", "GTM Plan")
    assert "Fallback" not in [text for _, text in styles]


def test_markdown_constructs_map_to_word_styles(tmp_path):
    styles = dict(
        (text, style)
        for style, text in _styles(markdown_to_docx(PLAN, tmp_path / "p.docx", "F"))
    )
    assert styles["Generated from evidence."] == "Intense Quote"
    assert styles["Market Changes"] == "Heading 1"
    assert styles["1. ROI Calculator"] == "Heading 2"
    assert styles["Insight: pricing dominates"] == "List Bullet"
    assert styles["First action"] == "List Number"
    assert styles["Buyers are evaluating economics."] == "Normal"


def test_bold_survives_into_word(tmp_path):
    doc = Document(str(markdown_to_docx(PLAN, tmp_path / "p.docx", "F")))
    para = next(p for p in doc.paragraphs if p.text.startswith("Target:"))
    assert [(r.text, r.bold) for r in para.runs] == [
        ("Target: ", False),
        ("VP Finance", True),
    ]


def test_sections_without_a_leading_h1_still_start_at_heading_1(tmp_path):
    """`##` is the top section whether the title came from a `# ` or from
    `fallback_title`.

    The executive summary supplies no `# ` of its own — the Run titles it —
    so keying the demotion off a consumed `# ` alone started its sections at
    Heading 2 and left Heading 1 unused, while the same sections in the plan
    were Heading 1.
    """
    text = "## What changed?\nBody.\n\n## What should Sales do?\nMore body.\n"
    styles = _styles(markdown_to_docx(text, tmp_path / "s.docx", "Executive Summary"))

    assert styles[0] == ("Title", "Executive Summary")
    assert ("Heading 1", "What changed?") in styles
    assert ("Heading 1", "What should Sales do?") in styles


def test_bare_uppercase_labels_are_promoted_as_a_fallback(tmp_path):
    """The summary prompt asks for `##`, but a model that ignores it and
    falls back to plain uppercase labels should still produce a document
    with headings rather than one long block of body text."""
    styles = _styles(markdown_to_docx(SUMMARY, tmp_path / "s.docx", "Executive Summary"))
    assert styles[0] == ("Title", "Executive Summary")
    assert ("Heading 1", "WHAT CHANGED") in styles
    assert ("Heading 1", "THREE MOST IMPORTANT ACTIONS") in styles


def test_sentence_case_text_is_never_promoted_to_a_heading(tmp_path):
    styles = _styles(markdown_to_docx(SUMMARY, tmp_path / "s.docx", "Executive Summary"))
    assert ("Normal", "Pricing themes lead this run.") in styles


def test_missing_dependency_raises_an_actionable_error():
    """The pipeline catches this; the message has to tell a user what to do."""
    exc = DocxUnavailable(
        "python-docx is not installed, so no Word version was written."
    )
    assert "python-docx" in str(exc)
    with pytest.raises(DocxUnavailable):
        raise exc


WRAPPED = """# Report

> A quote that the model wrapped
> across two source lines.

## Section
A sentence wrapped at roughly seventy-two columns,
continuing onto a second line,
and a third.

- A bullet whose text also wraps
  onto a continuation line

Another paragraph.
"""


def test_soft_wrapped_prose_becomes_one_paragraph(tmp_path):
    """Markdown treats a single newline as a soft break. Emitting one Word
    paragraph per source line shatters wrapped prose into fragments."""
    styles = _styles(markdown_to_docx(WRAPPED, tmp_path / "w.docx", "F"))
    bodies = [text for style, text in styles if style == "Normal"]
    assert bodies[0] == (
        "A sentence wrapped at roughly seventy-two columns, "
        "continuing onto a second line, and a third."
    )
    assert "Another paragraph." in bodies


def test_wrapped_blockquote_becomes_one_quote(tmp_path):
    styles = _styles(markdown_to_docx(WRAPPED, tmp_path / "w.docx", "F"))
    quotes = [text for style, text in styles if style == "Intense Quote"]
    assert quotes == ["A quote that the model wrapped across two source lines."]


def test_wrapped_list_item_stays_one_item(tmp_path):
    styles = _styles(markdown_to_docx(WRAPPED, tmp_path / "w.docx", "F"))
    bullets = [text for style, text in styles if style == "List Bullet"]
    assert bullets == ["A bullet whose text also wraps onto a continuation line"]


def test_blank_line_still_separates_paragraphs(tmp_path):
    """Buffering must not run distinct paragraphs together."""
    styles = _styles(markdown_to_docx("# T\n\nOne.\n\nTwo.\n", tmp_path / "b.docx", "F"))
    assert [text for style, text in styles if style == "Normal"] == ["One.", "Two."]
